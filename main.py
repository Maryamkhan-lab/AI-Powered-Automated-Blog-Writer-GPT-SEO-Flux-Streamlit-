import io, os, re, json, shutil, requests
import time
from PIL import Image
from io import BytesIO
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from tqdm import tqdm
from PIL import Image
from pprint import pprint
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from typing import List, Dict
from pydantic import BaseModel
from docx.shared import Inches
from docx.shared import RGBColor
from docx.oxml import OxmlElement
from fastapi import FastAPI, HTTPException ,Query
from fastapi.responses import FileResponse
from fastapi.responses import StreamingResponse
from langchain_core.prompts import PromptTemplate
from dotenv import load_dotenv
import logging
from fastapi import Request


load_dotenv()

app = FastAPI()

class ImagePromptRequest(BaseModel):
    blog_post_content: str
    previous_image_prompts: str

API_KEY = os.getenv("OPENAI_API_KEY")

from langchain_openai import ChatOpenAI

llm = ChatOpenAI(model="gpt-4o",
                temperature=1, 
                max_tokens=1024, 
                api_key=API_KEY
                )

class blog_request(BaseModel):
    TypeOf : str
    target_audience: str
    tone: str
    point_of_view: str
    target_country: str
    keywords: List[str]
    category: List[str]
    subheadings: int
    

def fetch_google_results(keywords: List[str], target_country: str) -> List[str]:
    username=os.getenv("USERNAME")
    password=os.getenv("PASSWORD")
    all_results_dict = {}
    for keyword in keywords:
        payload = {'source': 'google_search','query': keyword, 'domain': 'com', 'geo_location': target_country, 'locale': 'en-us', 'parse': True, 'start_page': 1, 'pages': 5,  'limit': 10, }
    try:
        response = requests.post('https://realtime.oxylabs.io/v1/queries', auth=(username, password), json=payload )
        response.raise_for_status()
        all_results_dict[keyword] = response.json()
    except requests.RequestException as e:
        raise HTTPException(status_code=500, detail=f"Error for '{keyword}': {str(e)}")
    formatted_results = {keyword: {'results': [{ 'pos': organic.get('pos'), 'url': organic.get('url'),'title': organic.get('title')}
                for result in all_results_dict[keyword].get('results', [])
                for organic in result.get('content', {}).get('results', {}).get('organic', [])   ]        }    }
    return formatted_results

def generate_blog_title(keywords: List[str], search_results: List[str],blog_request: blog_request) -> str:
    prompt_template = """
    You are an expert content creator and SEO strategist. Your task is to craft a single, SEO-optimized, and reader-focused title for a blog post using the provided keywords. 

    Blog Post Details:
    - **Category**: {category}
    - **Keywords**: {keywords}
    - **Type**: {Type}

    Instructions:
    1. Use the given keywords naturally and effectively within the title.
    2. Ensure the title is concise (preferably under 50 characters) while retaining clarity and relevance.
    3. Make it catchy and engaging to attract readers’ attention.
    4. Reflect the blog's type and purpose (e.g., listicle, guide, how-to, etc.) in the title, catering specifically to readers interested in {category}.
    5. Maintain a positive, inviting tone that aligns with the topic.
    6. Avoid generic or vague phrases; ensure the title is specific and impactful.

    Output:
    - Provide a single SEO-friendly title only, without additional explanations or formatting.
"""

    prompt = prompt_template.format(
        Type=blog_request.TypeOf,
        category=blog_request.category,
        keywords=", ".join(keywords),
        search_results="\n".join(search_results)
    )

    response = llm.invoke(prompt)
    
    return response.content

    
def generate_blog_subheadings(title: str, search_results : list, blog_request: blog_request) -> List[str]:
    prompt_template = """
    You are a skilled content strategist and SEO expert tasked with creating compelling and SEO-optimized subheadings for a blog post. These subheadings should enhance readability, engage the target audience, and align with the blog’s title and focus.

    Blog Post Details:
    - **Title**: {title}
    - **Category**: {seleted_catagory}
    - **SEO Keywords**: {search_results}
    - **Target Audience**: {target_audience}
    - **Tone**: {tone}
    - **Point of View**: {point_of_view}
    - **Target Country**: {target_country}

    Instructions:
    1. Generate {subheadings} subheadings that comprehensively address important aspects of the topic.
    2. Ensure each subheading incorporates relevant keywords and resonates with the blog's tone and target audience.
    3. Use concise, clear, and engaging language that encourages readers to continue exploring the blog.
    4. If applicable, include tips, actionable insights, or region-specific details to add value.
    5. Maintain a logical flow between subheadings to create a seamless reading experience.
    6. Write only subheadings, and do not include any additional text or formatting in the output.

    Based on this input, suggest SEO-friendly subheadings for the blog post.
"""

    prompt = prompt_template.format(
        title=title, seleted_catagory=blog_request.category ,search_results=", ".join(search_results), target_audience=blog_request.target_audience, 
        tone=blog_request.tone,point_of_view=blog_request.point_of_view,target_country=blog_request.target_country , subheadings=blog_request.subheadings
    )

    response = llm.invoke(prompt)
    
    suggested_subheadings = response.content.split("\n")
    return [subheading.strip() for subheading in suggested_subheadings if subheading.strip()]
    
def BlogPostPromptSingleSubheading(title: str, current_subheading: str, blog_request: blog_request, search_results: List[str], previous_content: str) -> str:
    prompt_template = """
    You are an expert content creator and language model specializing in crafting professional and engaging blog posts. 
    Your goal is to write a well-structured, SEO-optimized, and captivating section under the given subheading, 
    tailored to the target audience and aligned with the overall blog theme.

    Blog Post Details:
    - **Category**: {category}
    - **Title**: {title}
    - **Target Audience**: {target_audience}
    - **Tone**: {tone}
    - **Point of View**: {point_of_view}
    - **Target Country**: {target_country}
    
    Previous Content (for reference and context):
    {previous_content}
    
    **Subheading**: {current_subheading}

    Instructions:
    1. Write an engaging, unique, and factual section for the given subheading, ensuring it aligns seamlessly with the preceding content.
    2. Optimize the section for SEO by naturally integrating the provided keywords. Do not overuse them; maintain readability and flow.
    3. Use a tone that resonates with the target audience (e.g., reassuring, authoritative, or conversational) and fits the blog’s theme.
    4. Incorporate relevant research findings, statistics, expert quotes, or actionable advice to enrich the content and make it credible.
    5. Include practical tips, relatable examples, or insights that address the audience’s needs, questions, or challenges.
    6. Use smooth transitions to maintain consistency and guide readers into the subsequent sections effortlessly.
    7. Avoid irrelevant details, filler content, or generic phrases. Keep every sentence valuable and impactful.
    8. Use numbers for lists or tips (e.g., "1.", "2.", "3.") for clarity and organization. Do not use ** for lists or tips.
    9. For headings and subheadings, use ** (e.g., **Subheading**) to maintain consistent formatting.

    Additional Notes:
    - Each section should be concise yet comprehensive (maximum 2-3 paragraphs per subheading).
    - Focus on writing for readers first, with SEO considerations seamlessly integrated.
    - Avoid adding conclusions, references, or FAQs in the content.
    - Ensure the writing naturally leads into the next subheading.

    Now, based on this input, draft a compelling and SEO-friendly section for the given subheading.
"""

    prompt = prompt_template.format(title=title, category=blog_request.category,target_audience=blog_request.target_audience,tone=blog_request.tone,keywords=", ".join(search_results),point_of_view=blog_request.point_of_view,target_country=blog_request.target_country,previous_content=previous_content,current_subheading=current_subheading )

    response = llm.invoke(prompt)
    
    content=response.content
    return content
    
def format_content(document, content: str):
    subheading_pattern = r"\*\*(.*?)\*\*"
    sub_subheading_pattern = r"^(?:\d+\.)?\s*(.*?)\s*:\s*$"  # Pattern to detect subheading within subheading
    bullet_point_pattern = r"^\s*•\s*\*\*\s*(.*?)\s*\*\*"
    
    lines = content.split("\n")
    for line in lines:
        if re.match(subheading_pattern, line):
            # Main subheading
            subheading_text = re.sub(r"\*\*", "", line).strip()
            document.add_heading(subheading_text, level=2)
        elif re.match(sub_subheading_pattern, line):
            # Sub-subheading, promoted to a heading
            sub_subheading_text = re.match(sub_subheading_pattern, line).group(1).strip()
            document.add_heading(sub_subheading_text, level=3)
        elif re.match(bullet_point_pattern, line):
            # Bullet points
            bullet_text = re.sub(r"^\s*•\s*\*\*", "", line).strip()
            p = document.add_paragraph(style='List Bullet')
            run = p.add_run(bullet_text)
            run.bold = True
        else:
            # Regular paragraph
            p = document.add_paragraph(line.strip())
            p.alignment = 3  # Justify alignment
    
    # Set text color for all text in the document
    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            font = run.font
            font.color.rgb = RGBColor(0, 0, 0)  # Set text color to black

    
def generate_image_prompt(content: str, previous_prompts: str) -> str:
    prompt_template = """
You are a creative assistant tasked with generating visually stunning, realistic image prompts for a blog. Each prompt should be meticulously detailed, emotionally engaging, and tailored to enhance the blog's storytelling and themes.

Blog Post Draft:
{blog_post_content}

Previous Image Prompts:
{previous_image_prompts}

Instructions:
- Create image prompts that evoke emotional resonance, enhance storytelling, and maintain professional visual appeal.
- Include the following key components:
  1. **Subject/Scene**: Clearly specify the subject(s) or central themes. Highlight moments of connection, action, or visually engaging scenarios such as daily life, nature, or creative projects.
  2. **Composition and Action**: Describe spatial arrangements, depth, and active storytelling moments. For example, a person journaling in a cozy nook, a couple walking hand-in-hand in a scenic park, or an artist sketching in a studio.
  3. **Emotion and Style**: Convey a mood or artistic style (e.g., natural, cinematic, editorial). Emphasize emotional moments like joy, introspection, or vibrancy.
  4. **Lighting and Color**: Use lighting to enhance the tone, such as soft natural light, dramatic shadows, or vibrant hues. Specify color palettes that suit the scene, like earthy tones, bright pops of color, or monochromatic schemes.
  5. **Camera and Lens Settings**: Recommend camera models (e.g., Canon EOS R5, Sony Alpha 7R IV), lenses (e.g., 35mm f/1.8 for general scenes or 85mm for portraits), and techniques (e.g., shallow depth of field, long exposure for motion).
  6. **Artistic Enhancements**: Suggest details like angles (e.g., bird’s-eye view, close-up), effects (e.g., bokeh, motion blur), or scene accents (e.g., props, textures, or natural elements).
  7. **Aspect Ratio and Style Tags**: Specify dimensions (e.g., --ar 16:9 for banners or --ar 4:5 for Instagram). Include style tags like --style cinematic, --style raw, or --style editorial.

Examples:
1. A solitary hiker standing at the edge of a cliff overlooking a vast mountain range at sunrise. Warm, golden light bathes the scene, and a shallow depth of field highlights the hiker while keeping the horizon slightly blurred. Captured with a Sony Alpha 1 and a 24-70mm lens. --ar 16:9 --style cinematic

2. A cozy living room scene featuring a person relaxing on a sofa with a steaming mug of tea, surrounded by soft blankets and fairy lights. Gentle warm light from a nearby lamp enhances the feeling of comfort. Taken with a Canon EOS R6 and a 50mm f/1.4 lens. --ar 4:5 --style editorial

3. A bustling city street during golden hour, capturing pedestrians, cyclists, and the glow of sunlight reflecting off glass buildings. The composition emphasizes urban energy, with a wide-angle shot creating depth. Shot with a Nikon Z9 and a 24mm lens. --ar 16:9 --style candid

4. A flat-lay shot of an artist’s workspace, showcasing paintbrushes, vibrant palettes, and an unfinished canvas. The scene is styled with natural textures like wood and linen, and diffused sunlight streams through a nearby window. Shot with a Fujifilm X-T4 and a 35mm f/2 lens. --ar 1:1 --style clean

5. A serene close-up of dewdrops on a leaf at dawn, with soft focus highlighting the water droplets and delicate leaf veins. Captured in natural light using a macro lens for fine detail. Shot with a Canon EOS R5 and a 100mm macro lens. --ar 3:2 --style raw

"""

    prompt = prompt_template.format(blog_post_content=content,previous_image_prompts=previous_prompts)

    response = llm.invoke(prompt)
    
    return response.content  # Extract prompts


# Dictionary to store blog content and images
storage = {
    "blog_post": None,
    "images": {}
}

def generate_image(prompt: str):
    IMAGE_API_KEY = os.getenv("IMAGE_API_KEY")
    url = "https://api.bfl.ml/v1/flux-pro-1.1"
    headers = {
        "accept": "application/json",
        "x-key": IMAGE_API_KEY,
        "Content-Type": "application/json"
    }
    payload = {
        "prompt": prompt,
        "width": 1024,
        "height": 1024,
        "guidance_scale": 1,
        "num_inference_steps": 50,
        "max_sequence_length": 512,
        'Safety Tolerance': 3,
    }
    
    # Sending the initial request to generate the image
    response = requests.post(url, headers=headers, json=payload).json()
    if "id" not in response:
        print("Error generating image:", response)
        return None
    
    request_id = response["id"]
    
    # Polling for the result
    while True:
        time.sleep(0.5)
        result = requests.get(
            "https://api.bfl.ml/v1/get_result",
            headers=headers,
            params={"id": request_id},
        ).json()
        
        # Check the status of the result
        status = result.get("status")
        if status == "Ready":
            if "result" in result and "sample" in result["result"]:
                image_url = result["result"]["sample"]
                image_response = requests.get(image_url)
                if image_response.status_code == 200:
                    image = Image.open(BytesIO(image_response.content))
                    return image
            else:
                print("Error: No 'sample' key in result.")
                return None
        elif status == "Content Moderated":
            print("Image generation status: Content Moderated. Stopping generation.")
            break
        else:
            print(f"Image generation status: {status}")

def selected_category(category: dict, search_results: list) -> str:
    prompt_template = """
    Based on the given search results, select the most appropriate category for the blog post.
    Available Categories: {categories}
    Search Results: 
    {search_results}
    Carefully analyze the keywords and context in the search results to choose the best category. 
    Please respond only with the most relevant category name.
    """
    prompt = prompt_template.format(categories=", ".join(category.keys()), search_results="\n".join(search_results))

    response = llm.invoke(prompt)
    
    return response.content.strip()  # Extract the selected category
    
def fetch_google_results_for_site(keywords: List[str]) -> List[Dict[str, int]]:
    USERNAME = os.getenv("USERNAME")
    PASSWORD = os.getenv("PASSWORD")
    query_string = "+".join(keywords)
    search_url = f"https://www.google.com/search?q=site:marcusmcdonnell.com+{query_string}"
    payload = {
        'source': 'google',
        'url': search_url,
        'parse': True  # Enabling parsed response to get structured data
    }
    try:
        response = requests.post(
            'https://realtime.oxylabs.io/v1/queries',
            auth=(USERNAME, PASSWORD),
            json=payload
        )
        response.raise_for_status()
        full_response = response.json()
        filtered_results = []
        if full_response.get('results'):
            for result in full_response['results']:
                organic_results = result.get('content', {}).get('results', {}).get('organic', [])
                if isinstance(organic_results, list):
                    filtered_results.extend(
                        {"title": item.get("title"), "url": item.get("url"), "pos": item.get("pos")}
                        for item in organic_results
                        if "title" in item and "url" in item and "pos" in item                    )
                else:
                    print("Expected 'organic' results to be a list but found something else.")
        else:
            print("No 'results' key found in the response.")
        return filtered_results
    except requests.RequestException as e:
        print(f"Error fetching results: {e}")
        return []
        
def generate_linkages(blog_post: str, search_results: list, keywords: List[str]) -> dict:
    Internal_search_results = fetch_google_results_for_site(keywords)
    prompt_template = """
    Based on the given blog post and search results, generate relevant external and internal links.
    
    Blog Post:
    {blog_post}
    
    Use the top 3 search results for external link suggestions, considering their relevance and quality. The links should be clickable hyperlinks.
    Also, suggest internal links that might help the reader based on the blog post's content. Do not include placeholder statements like 'no links found.'
    
    External Links:
    Provide a list of up to 3 high-quality external links with a brief description of each link's content and its relevance to the blog post. Ensure all links are clickable.
    
    Internal Links:
    Suggest up to 3 internal links based on the blog post's content. Provide a brief explanation of how each internal link connects to the blog post.
    
    External Links Results: 
    {search_results}
    
    Internal Links Results:
    {Internal_search_results}
    
    Output:
    External Links: 
    - [Link Text](URL): Brief explanation of relevance.
    
    Internal Links:
    - [Link Text](URL): Brief explanation of relevance.
    """

    prompt = prompt_template.format(blog_post=blog_post, search_results=search_results, Internal_search_results=Internal_search_results)

    response = llm.invoke(prompt)
    
    result = response.content.strip()  # Adjust based on LLM output structure
    
    return result
    
import io
from fastapi import FastAPI, HTTPException
from fastapi.responses import StreamingResponse
from docx import Document
from tqdm import tqdm
from typing import Dict
import os
import shutil


@app.post("/generate_blog/", response_model=dict)
def create_blog_pipeline(blog_request: blog_request):
    try:
               
        print('SEO Searching')
        search_results = fetch_google_results(blog_request.keywords, blog_request.target_country)

        print('Generating Title for blog post')
        previous_image_prompts = ''
        blog_content = ""
        document = Document()

        title = generate_blog_title(blog_request.keywords, search_results, blog_request)
        document.add_heading(title, 0)

        image_prompt = generate_image_prompt(title, previous_image_prompts)
        previous_image_prompts += image_prompt + " , "
        image = generate_image(image_prompt)

        if image:
            image_stream = io.BytesIO()
            image.save(image_stream, format="PNG")
            image_stream.seek(0)
            document.add_picture(image_stream, width=Inches(6), height=Inches(6))
        else:
            print("Title image generation failed.")

        print('Generating Subheadings for blog post')
        subheadings = generate_blog_subheadings(title, search_results, blog_request)

        for i, subheading in enumerate(tqdm(subheadings, desc="Processing subheadings")):
            content = BlogPostPromptSingleSubheading(
                title, subheading, blog_request, search_results, blog_content
            )
            blog_content += f"\n\n{subheading}\n{content}"
            format_content(document, content)

            image_prompt = generate_image_prompt(content, previous_image_prompts)
            previous_image_prompts += image_prompt + " , "
            image = generate_image(image_prompt)
            if image:
                image_stream = io.BytesIO()
                image.save(image_stream, format="PNG")
                image_stream.seek(0)
                document.add_picture(image_stream, width=Inches(6), height=Inches(6))
            else:
                print(f"Image generation failed for subheading: {subheading}")

        # Generate linkages and format them
        raw_linkages = generate_linkages(blog_content, search_results, blog_request.keywords)
        document.add_heading("Relevant Links", level=2)
        formatted_linkages = format_content(document, raw_linkages)
        document.add_paragraph(formatted_linkages)
        

        # Save document to memory stream
        docx_stream = io.BytesIO()
        document.save(docx_stream)
        docx_stream.seek(0)
        storage["blog_post"] = docx_stream.getvalue()
        storage["title"] = title

        return {"message": "Blog post generated successfully", "title": title}
            
    except Exception as e:
        print(f"An error occurred: {e}")
        return {"error": str(e)}
    
    except Exception as e:
        print(f"An error occurred: {e}")
        return {"error": str(e)}

@app.get("/download/")
def download_file():
    if not storage["blog_post"]:
        raise HTTPException(status_code=404, detail="No blog post available for download.")
    
    blog_title = storage.get("title", "Generated_Blog_Post")

    return StreamingResponse(
        io.BytesIO(storage["blog_post"]),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{blog_title}.docx"'}
    )


@app.get("/")
async def root():
    return {"message": "API is up and running!"}

@app.middleware("http")
async def log_requests(request: Request, call_next):
    print(f"Incoming request: {request.method} {request.url}")
    response = await call_next(request)
    print(f"Response status: {response.status_code}")
    return response